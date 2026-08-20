import html
import logging
import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from markdown_it import MarkdownIt

from app.services.book_parser import BookParser
from app.services.source_text import read_source_text
from app.services.universal_parser import UniversalParser
from app.types import ParsedDocument, Section

# HTML tag stripper for EPUB content
_RE_HTML_TAGS = re.compile(r"<[^>]+>")
_RE_WHITESPACE = re.compile(r"\s+")

# An EPUB may pack many chapters into one document, so sections come from
# headings rather than files.
_RE_EPUB_HEADING = re.compile(r"<h([1-6])[^>]*>(.*?)</h\1>", re.IGNORECASE | re.DOTALL)
# Block ends become paragraph breaks before tags are stripped; collapsing all
# whitespace first would leave each chapter a single run-on line.
_RE_EPUB_BLOCK_END = re.compile(
    r"</(?:p|div|h[1-6]|li|blockquote|tr|section|article)\s*>|<br\s*/?>",
    re.IGNORECASE,
)
_RE_BLANK_RUN = re.compile(r"\n{3,}")
_RE_INLINE_SPACE = re.compile(r"[ \t\r\f\v]+")


def _epub_text(fragment: str) -> str:
    """HTML fragment to reading text, with paragraph breaks preserved."""
    text = _RE_EPUB_BLOCK_END.sub("\n\n", fragment)
    text = _RE_HTML_TAGS.sub(" ", text)
    text = html.unescape(text)
    text = _RE_INLINE_SPACE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _RE_BLANK_RUN.sub("\n\n", text).strip()


def _split_epub_document(raw_html: str, fallback_heading: str) -> list[tuple[str, str]]:
    """Split one EPUB document into (heading, text) per heading tag it contains.

    Text before the first heading is kept under `fallback_heading` so front
    matter is not dropped. A document with no headings yields a single entry.
    """
    matches = list(_RE_EPUB_HEADING.finditer(raw_html))
    if not matches:
        return [(fallback_heading, _epub_text(raw_html))]

    out: list[tuple[str, str]] = []
    preamble = _epub_text(raw_html[: matches[0].start()])
    if preamble:
        out.append((fallback_heading, preamble))

    for i, m in enumerate(matches):
        heading = _epub_text(m.group(2)) or fallback_heading
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_html)
        out.append((heading, _epub_text(raw_html[m.end() : end])))
    return out

# Kindle clippings separator
_KINDLE_SEP = "=========="
# Highlight/note header: "- Your Highlight on page X | Added on Date"
_RE_KINDLE_HEADER = re.compile(
    r"^-\s+Your\s+(Highlight|Note|Bookmark)\s+.*\|\s+Added\s+on\s+(.+)$",
    re.IGNORECASE,
)

logger = logging.getLogger(__name__)

# Shared instances (stateless)
_book_parser = BookParser()
_universal_parser = UniversalParser()

# "Chapter 5 describes...", "Part 3 are...", "volume 22 of..." — a chapter/part
# marker followed by a lowercase word is prose, not a heading. Real headings
# capitalize the title ("Chapter 5 Machine Learning") or use punctuation.
_RE_MARKER_PROSE = re.compile(
    r"^(?:chapter|chap|part|book|section|volume|adventure)\s+\w+\s+(\w+)",
    re.IGNORECASE,
)


def _norm_ws(s: str) -> str:
    """Collapse the multi-space runs PyMuPDF leaves between spans."""
    return _RE_WHITESPACE.sub(" ", s).strip()


# Google Docs PDF exports name every bookmark after its HTML anchor, so the
# TOC reads `_r9szt46p8rxa` instead of the heading the reader sees.
_RE_ANCHOR_ID = re.compile(r"^_[a-z0-9]{5,}$", re.IGNORECASE)
_DERIVED_HEADING_MAX = 80


def _has_words(text: str) -> bool:
    letters = sum(ch.isalpha() for ch in text)
    return letters >= 2 and letters >= len(text) * 0.5


def _heading_from_page(page, body_size: float) -> str:
    """The largest heading-sized line on a page, or "".

    The bookmark's destination is trustworthy even when its title is not, so the
    real heading is on the page it points at, set larger than body text. Reading
    the section's opening text instead returns the running header, which is the
    same on every page of a chapter.
    """
    best_size = 0.0
    best_text = ""
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            text = _norm_ws(_join_spans(spans))
            if not text or len(text) > _DERIVED_HEADING_MAX:
                continue
            # A continuation page has no heading, and its largest line is the
            # page number. Require real words.
            if not _has_words(text):
                continue
            size = max((s.get("size", 0.0) for s in spans), default=0.0)
            if size > best_size:
                best_size, best_text = size, text
    return best_text if best_size > body_size * 1.1 else ""


def _usable_heading(title: str, page, body_size: float, fallback_body: str) -> str:
    """The bookmark title, or the heading actually printed at its destination."""
    clean = _norm_ws(title)
    if clean and not _RE_ANCHOR_ID.match(clean):
        return clean
    recovered = _heading_from_page(page, body_size)
    if recovered:
        return recovered
    first_line = next(
        (ln.strip() for ln in fallback_body.splitlines() if _has_words(ln.strip())), ""
    )
    return _norm_ws(first_line)[:_DERIVED_HEADING_MAX].strip() or clean


# Kerning within a word is ~0 (often negative); a space glyph is 0.25-0.33em.
_SPACE_GAP_EM = 0.2


# Running headers and footers sit in the page margins and are short. Once
# blocks become paragraphs, each would otherwise land in the prose as one.
_MARGIN_FRACTION = 0.1
_FURNITURE_MAX_CHARS = 80


def _printed_page_labels(doc) -> dict[int, str]:
    """Sheet number -> the number printed on it, where the two differ.

    A book numbers its front matter separately, so a PDF carries page *labels*
    beside sheet positions: measured on a 613-page book, sheet 41 is printed
    "19" and sheet 6 is printed "iv". A citation naming the sheet therefore
    disagrees with the page in the reader's hands by twenty for the whole body.

    Only differing entries are kept: on three of four books in one library the
    PDF defines no labels at all, and storing "5" for sheet 5 would be noise
    that later code has to re-derive nothing from.
    """
    labels: dict[int, str] = {}
    try:
        for index in range(doc.page_count):
            raw = (doc[index].get_label() or "").strip()
            if raw and raw != str(index + 1):
                labels[index + 1] = raw
    except Exception:  # noqa: BLE001 - labels are a nicety; parsing must not fail for them
        logger.warning("PDF page labels unavailable; citations will name sheet numbers")
        return {}
    return labels


def _line_start_offsets(lines: list[str]) -> list[int]:
    """Character offset at which each line starts once joined by a newline."""
    offsets: list[int] = []
    running = 0
    for line in lines:
        offsets.append(running)
        running += len(line) + 1  # the "\n" the lines are joined with
    return offsets


def _block_start_offsets(blocks: list[str]) -> list[int]:
    """Character offset at which each block starts once joined by a blank line."""
    offsets: list[int] = []
    running = 0
    for block in blocks:
        offsets.append(running)
        running += len(block) + 2  # the "\n\n" the blocks are joined with
    return offsets


def _is_page_furniture(block: dict, page_height: float) -> bool:
    bbox = block.get("bbox")
    if not bbox or len(bbox) < 4:
        return False
    top, bottom = float(bbox[1]), float(bbox[3])
    margin = page_height * _MARGIN_FRACTION
    in_margin = bottom <= margin or top >= page_height - margin
    if not in_margin:
        return False
    text = " ".join(
        _join_spans(line.get("spans", [])) for line in block.get("lines", [])
    ).strip()
    return 0 < len(text) <= _FURNITURE_MAX_CHARS


def _join_spans(spans: list[dict]) -> str:
    """Rebuild a line's text from its spans, inferring word gaps from geometry.

    PyMuPDF opens a new span at every style change, so a PDF alternating Type3
    font subsets per glyph splits words across dozens of spans. Joining on a
    space corrupts every word; joining on nothing welds together words a PDF
    separates by position alone. Only the gap tells the two apart.
    """
    parts: list[str] = []
    prev_x1: float | None = None
    prev_text = ""
    for span in spans:
        text = span.get("text", "")
        if not text:
            continue
        x0, x1 = span["bbox"][0], span["bbox"][2]
        if (
            prev_x1 is not None
            and not prev_text[-1:].isspace()
            and not text[:1].isspace()
            and x0 - prev_x1 > _SPACE_GAP_EM * span.get("size", 12.0)
        ):
            parts.append(" ")
        parts.append(text)
        prev_x1 = x1
        prev_text = text
    return "".join(parts)


def _heading_is_prose(heading: str) -> bool:
    h = heading.strip()
    if len(h) > 100:
        return True
    m = _RE_MARKER_PROSE.match(h)
    return bool(m and m.group(1)[:1].islower())


def _sections_plausible(sections: list[Section]) -> bool:
    """Reject a segmentation whose headings are mostly prose fragments."""
    if not sections:
        return False
    prose = sum(1 for s in sections if _heading_is_prose(s.heading))
    return prose < len(sections) / 2


class DocumentParser:
    """
    General document parser.

    Uses a tiered approach:
    1. UniversalParser (signature discovery for books, scripts, papers, chat)
    2. BookParser (legacy regex families for classic books)
    3. Heuristic fallbacks (font-size, paragraph splits)
    """

    def parse(self, file_path: Path, format: str) -> ParsedDocument:
        fmt = format.lower()
        if fmt == "pdf":
            return self._parse_pdf(file_path)
        elif fmt == "docx":
            return self._parse_docx(file_path)
        elif fmt == "txt":
            return self._parse_txt(file_path)
        elif fmt in ("md", "markdown"):
            return self._parse_md(file_path)
        elif fmt == "epub":
            return self._parse_epub(file_path)
        else:
            raise ValueError(f"Unsupported format: {format}")

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        # Guard: fitz can hang (not raise) on garbage bytes; fail fast here.
        with file_path.open("rb") as _f:
            if not _f.read(5).startswith(b"%PDF-"):
                raise ValueError(f"Not a valid PDF (missing %PDF- header): {file_path.name}")

        # If the PDF ships with an embedded TOC, prefer the TOC path below:
        # it reads page numbers directly from PyMuPDF's resolved bookmarks,
        # which is more reliable than BookParser's substring-search heuristic
        # (that fails on titles split across lines or rendered in different
        # case than the heading string, e.g. "CHAPTER 4\nOptimizing..." vs
        # "Chapter 4. Optimizing...", causing 1-page-off matches via running
        # headers on the next page).
        doc = fitz.open(str(file_path))

        # Skip UniversalParser for PDFs -- the font-size heuristic below
        # leverages actual font metrics from the PDF structure and produces
        # better section boundaries than regex-based signature discovery.
        # UniversalParser is designed for plain text where font info is absent.
        total_pages = len(doc)
        # Read once here so every return below carries them, including the
        # fallback paths -- two of four books in one library take those.
        page_labels = _printed_page_labels(doc)

        all_font_sizes: list[float] = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:  # type: ignore[arg-type]
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            all_font_sizes.append(span["size"])

        body_avg = sum(all_font_sizes) / len(all_font_sizes) if all_font_sizes else 12.0
        heading_threshold = body_avg * 1.2

        # Try PDF built-in bookmarks (TOC) first -- these have accurate page numbers
        # and define the chapter-level structure the author intended.
        toc = doc.get_toc()  # [[level, title, page_num (1-based)], ...]

        if toc:
            # The entries themselves are not logged: a 1,017-entry manual wrote a
            # single INFO record thousands of lines long on every ingest, which
            # buries everything else in the log it shares.
            logger.info("PDF TOC: %d entries", len(toc))
            logger.debug("PDF TOC entries: %s", toc)
            # Use the full TOC hierarchy to build sections directly.
            # Trust the TOC structure -- it reflects what the author intended.
            # Font-based sub-heading detection is skipped here to avoid
            # mis-classifying bold/emphasis text as section boundaries in
            # properly structured PDFs.
            sections: list[Section] = []
            raw_parts: list[str] = []

            for i, (lv, ti, pg) in enumerate(toc):
                # End page = start of the next entry at the same or higher level,
                # minus 1. This correctly handles nested entries.
                next_page = total_pages + 1
                for j in range(i + 1, len(toc)):
                    if toc[j][0] <= lv:
                        next_page = toc[j][2]
                        break
                page_end = min(next_page - 1, total_pages)

                texts: list[str] = []
                # Index into `texts` at which each page after the first starts,
                # so a per-chunk page can be recovered later. Without it every
                # chunk in the section can only claim the section's start page.
                page_start_blocks: list[int] = []

                for page_offset, pn in enumerate(range(max(0, pg - 1), page_end)):
                    if page_offset > 0:
                        page_start_blocks.append(len(texts))
                    page_obj = doc[pn]
                    page_height = float(page_obj.rect.height) or 1.0
                    for block in page_obj.get_text("dict")["blocks"]:  # type: ignore[arg-type]
                        if block.get("type") != 0:
                            continue
                        if _is_page_furniture(block, page_height):
                            continue
                        block_lines: list[str] = []
                        for line in block.get("lines", []):
                            spans = line.get("spans", [])
                            if not spans:
                                continue
                            line_text = _join_spans(spans).strip()
                            if not line_text:
                                continue
                            block_lines.append(line_text)
                        if block_lines:
                            texts.append("\n".join(block_lines))

                # Blocks are the layout's own paragraphs; a PDF text layer is
                # hard-wrapped, so joining lines alone leaves no boundary and
                # the reader renders a whole chapter as one block.
                joined = "\n\n".join(texts)
                text = joined.strip()
                # Offsets are measured on the joined string, then shifted by
                # whatever the strip removed from the front, so they stay true to
                # the text that is actually stored.
                lead = len(joined) - len(joined.lstrip())
                block_offsets = _block_start_offsets(texts)
                page_breaks = [
                    max(0, block_offsets[i] - lead)
                    for i in page_start_blocks
                    if i < len(block_offsets)
                ]
                raw_parts.append(text)
                sections.append(
                    Section(
                        heading=_usable_heading(
                            ti, doc[max(0, pg - 1)], body_avg, text
                        ),
                        level=lv,
                        text=text,
                        page_start=pg,
                        page_end=page_end,
                        page_breaks=page_breaks,
                    )
                )

            if sections:
                raw_text = "\n".join(raw_parts)
                return ParsedDocument(
                    title=file_path.stem,
                    format="pdf",
                    pages=total_pages,
                    word_count=len(raw_text.split()),
                    sections=sections,
                    raw_text=raw_text,
                    page_labels=page_labels,
                )

        # No TOC available -- single-pass font-size scan preserving document order.
        # Pre-scan all font sizes to find the two heading thresholds dynamically.
        all_sizes: list[float] = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:  # type: ignore[arg-type]
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if spans:
                        all_sizes.append(max(s["size"] for s in spans))

        # Distinct sizes above body threshold, sorted largest first.
        heading_sizes = sorted({s for s in all_sizes if s >= heading_threshold}, reverse=True)
        # Level-1 = the top distinct size group; level-2 = the rest above threshold.
        h1_min = heading_sizes[0] if heading_sizes else heading_threshold
        logger.info(
            "PDF fallback scan: body_avg=%.1f threshold=%.1f h1_min=%.1f distinct_heading_sizes=%s",
            body_avg,
            heading_threshold,
            h1_min,
            heading_sizes[:10],
        )

        sections = []
        raw_parts = []
        current_heading = "Introduction"
        current_level = 1
        current_page_start = 1
        current_texts: list[str] = []
        # Index into `current_texts` at which each later page of the current
        # section begins. One entry per page even when a page contributes no
        # lines, so counting entries below a position always yields the right
        # page across a blank leaf between chapters.
        current_page_marks: list[int] = []

        def flush_section(next_heading: str, next_level: int, next_page: int) -> None:
            nonlocal current_heading, current_level, current_page_start, current_texts
            nonlocal current_page_marks
            joined = "\n".join(current_texts)
            text = joined.strip()
            if text:
                lead = len(joined) - len(joined.lstrip())
                line_offsets = _line_start_offsets(current_texts)
                sections.append(
                    Section(
                        heading=_norm_ws(current_heading),
                        level=current_level,
                        text=text,
                        page_start=current_page_start,
                        page_end=next_page - 1,
                        page_breaks=[
                            max(0, line_offsets[i] - lead)
                            for i in current_page_marks
                            if i < len(line_offsets)
                        ],
                    )
                )
            current_heading = next_heading
            current_level = next_level
            current_page_start = next_page
            current_texts = []
            current_page_marks = []

        for page_num, page in enumerate(doc):
            if page_num + 1 > current_page_start:
                current_page_marks.append(len(current_texts))
            for block in page.get_text("dict")["blocks"]:  # type: ignore[arg-type]
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    max_size = max(s["size"] for s in spans)
                    line_text = _join_spans(spans).strip()
                    if not line_text:
                        continue
                    if max_size >= heading_threshold and len(line_text) < 120:
                        level = 1 if max_size >= h1_min else 2
                        flush_section(line_text, level, page_num + 1)
                    else:
                        current_texts.append(line_text)
                        raw_parts.append(line_text)

        flush_section("_end", 0, total_pages + 1)
        raw_text = "\n".join(raw_parts)

        # Font metrics are the structural signal for PDFs. Fall back to the
        # chapter-regex BookParser only when the font scan finds no plausible
        # structure (e.g. flat-font PDFs) -- and reject its result if it looks
        # like prose fragments, which is how a book that enumerates its own
        # chapters ("Chapter 5 describes...") used to get mis-segmented.
        if _sections_plausible(sections):
            return ParsedDocument(
                title=file_path.stem,
                format="pdf",
                pages=total_pages,
                word_count=len(raw_text.split()),
                sections=sections,
                raw_text=raw_text,
                page_labels=page_labels,
            )

        doc.close()
        book = _book_parser.parse(file_path, "pdf")
        if book is not None and _sections_plausible(book.sections):
            return book

        if not sections:
            sections = [
                Section(
                    heading="Introduction",
                    level=1,
                    text=raw_text,
                    page_start=1,
                    page_end=total_pages,
                )
            ]
        return ParsedDocument(
            title=file_path.stem,
            format="pdf",
            pages=total_pages,
            word_count=len(raw_text.split()),
            sections=sections,
            raw_text=raw_text,
            page_labels=page_labels,
        )

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        # Try UniversalParser first
        result = _universal_parser.parse(file_path, "docx")
        if result is not None:
            return result

        # Try BookParser next
        result = _book_parser.parse(file_path, "docx")
        if result is not None:
            return result

        # Fallback: Word heading styles only
        doc = DocxDocument(str(file_path))
        sections: list[Section] = []
        raw_parts: list[str] = []
        current_heading = "Introduction"
        current_level = 1
        current_texts: list[str] = []
        section_order = 0

        def flush_section(next_heading: str, next_level: int) -> None:
            nonlocal current_heading, current_level, current_texts, section_order
            text = "\n".join(current_texts).strip()
            if text:
                sections.append(
                    Section(
                        heading=current_heading,
                        level=current_level,
                        text=text,
                        page_start=0,
                        page_end=0,
                    )
                )
                section_order += 1
            current_heading = next_heading
            current_level = next_level
            current_texts = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if re.match(r"Heading\s+([1-3])", style_name):
                match = re.match(r"Heading\s+(\d+)", style_name)
                level = int(match.group(1)) if match else 1
                flush_section(text, level)
            else:
                current_texts.append(text)
                raw_parts.append(text)

        flush_section("_end", 0)

        raw_text = "\n".join(raw_parts)
        word_count = len(raw_text.split())
        title = file_path.stem
        return ParsedDocument(
            title=title,
            format="docx",
            pages=0,
            word_count=word_count,
            sections=sections,
            raw_text=raw_text,
        )

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def _parse_txt(self, file_path: Path) -> ParsedDocument:
        # Try BookParser first (legacy regex families for classic books)
        result = _book_parser.parse(file_path, "txt")
        if result is not None:
            return result

        # Try UniversalParser next (signature discovery)
        result = _universal_parser.parse(file_path, "txt")
        if result is not None:
            return result

        # Fallback: paragraph-split (for unstructured text files)
        text = read_source_text(file_path)
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        sections = [
            Section(
                heading=f"Section {i + 1}",
                level=1,
                text=para,
                page_start=0,
                page_end=0,
            )
            for i, para in enumerate(paragraphs)
        ]
        word_count = len(text.split())
        title = file_path.stem
        return ParsedDocument(
            title=title,
            format="txt",
            pages=0,
            word_count=word_count,
            sections=sections,
            raw_text=text,
        )

    # ------------------------------------------------------------------
    # Markdown
    # ------------------------------------------------------------------

    def _parse_md(self, file_path: Path) -> ParsedDocument:
        # Try UniversalParser first
        result = _universal_parser.parse(file_path, "md")
        if result is not None:
            return result

        # Try BookParser first (markdown-it token-based with min-chapter guard)
        result = _book_parser.parse(file_path, "md")
        if result is not None:
            return result

        # Fallback: same markdown-it logic without the chapter minimum guard
        text = file_path.read_text(encoding="utf-8")
        md = MarkdownIt()
        tokens = md.parse(text)

        sections: list[Section] = []
        current_heading = "Introduction"
        current_level = 1
        current_texts: list[str] = []
        in_heading = False
        pending_heading = ""
        pending_level = 1

        def flush_section(next_heading: str, next_level: int) -> None:
            nonlocal current_heading, current_level, current_texts
            body = "\n".join(current_texts).strip()
            if body:
                sections.append(
                    Section(
                        heading=current_heading,
                        level=current_level,
                        text=body,
                        page_start=0,
                        page_end=0,
                    )
                )
            current_heading = next_heading
            current_level = next_level
            current_texts = []

        for token in tokens:
            if token.type == "heading_open":
                in_heading = True
                pending_level = int(token.tag[1]) if token.tag else 1
            elif token.type == "inline" and in_heading:
                pending_heading = token.content
            elif token.type == "heading_close":
                in_heading = False
                flush_section(pending_heading, pending_level)
            elif token.type == "inline":
                current_texts.append(token.content)

        flush_section("_end", 0)

        word_count = len(text.split())
        title = file_path.stem
        return ParsedDocument(
            title=title,
            format="md",
            pages=0,
            word_count=word_count,
            sections=sections,
            raw_text=text,
        )

    # ------------------------------------------------------------------
    # EPUB
    # ------------------------------------------------------------------

    def _parse_epub(self, file_path: Path) -> ParsedDocument:
        """Extract chapters from an EPUB file using ebooklib."""
        import ebooklib  # noqa: PLC0415
        from ebooklib import epub  # noqa: PLC0415

        book = epub.read_epub(str(file_path), options={"ignore_ncx": True})

        # Derive title from metadata or filename
        title_meta = book.get_metadata("DC", "title")
        title = title_meta[0][0] if title_meta else file_path.stem.replace("_", " ").title()

        sections: list[Section] = []
        raw_parts: list[str] = []

        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            item_name = item.get_name()
            # Skip navigation/toc documents
            if any(k in item_name.lower() for k in ("nav", "toc", "ncx")):
                continue

            raw_html = item.get_content().decode("utf-8", errors="replace")
            fallback = Path(item_name).stem.replace("_", " ").replace("-", " ").title()
            for heading_text, plain in _split_epub_document(raw_html, fallback):
                if len(plain) < 50:
                    continue
                sections.append(
                    Section(
                        heading=heading_text,
                        level=1,
                        text=plain,
                        page_start=0,
                        page_end=0,
                    )
                )
                raw_parts.append(plain)

        if not sections:
            # Fallback: treat entire content as one section
            logger.warning("_parse_epub: no usable chapters found in %s", file_path)
            full_text = " ".join(raw_parts) if raw_parts else ""
            sections = [
                Section(
                    heading="Content",
                    level=1,
                    text=full_text or "(empty)",
                    page_start=0,
                    page_end=0,
                )
            ]

        raw_text = "\n\n".join(raw_parts)
        word_count = len(raw_text.split())
        logger.info(
            "_parse_epub: %d sections, %d words from %s",
            len(sections),
            word_count,
            file_path,
        )
        return ParsedDocument(
            title=title,
            format="epub",
            pages=0,
            word_count=word_count,
            sections=sections,
            raw_text=raw_text,
        )

    # ------------------------------------------------------------------
    # Kindle My Clippings.txt
    # ------------------------------------------------------------------

    @staticmethod
    def parse_kindle_clippings(text: str) -> list[ParsedDocument]:
        """Parse a Kindle My Clippings.txt file into one ParsedDocument per book.

        Returns a list of ParsedDocument objects, one per book title found.
        Each section in the document is one highlight (with date as heading).
        """
        # Split on the separator line (lines starting with ==========)
        entries = re.split(r"^==========[ \t]*$", text, flags=re.MULTILINE)

        # Group highlights by book title (first non-empty line of each entry)
        books: dict[str, list[Section]] = {}
        book_order: list[str] = []

        for entry in entries:
            lines = [ln.strip() for ln in entry.strip().splitlines() if ln.strip()]
            if len(lines) < 2:
                continue

            book_title = lines[0]
            # Strip any author in parentheses: "Title (Author Name)" -> "Title"
            book_title_clean = re.sub(r"\s*\([^)]*\)\s*$", "", book_title).strip()
            if not book_title_clean:
                continue

            # Find the metadata header line (contains "Your Highlight" or "Your Note")
            metadata_line = ""
            content_lines: list[str] = []
            header_found = False
            for line in lines[1:]:
                m = _RE_KINDLE_HEADER.match(line)
                if m and not header_found:
                    metadata_line = line
                    header_found = True
                elif header_found:
                    content_lines.append(line)

            if not content_lines:
                continue

            highlight_text = " ".join(content_lines).strip()
            if not highlight_text:
                continue

            # Use "Added on <date>" as the section heading
            date_str = ""
            date_match = re.search(r"Added on (.+)$", metadata_line)
            if date_match:
                date_str = date_match.group(1).strip()

            heading = f"Highlight ({date_str})" if date_str else "Highlight"

            if book_title_clean not in books:
                books[book_title_clean] = []
                book_order.append(book_title_clean)

            books[book_title_clean].append(
                Section(heading=heading, level=1, text=highlight_text, page_start=0, page_end=0)
            )

        documents: list[ParsedDocument] = []
        for book_title in book_order:
            sections = books[book_title]
            raw_text = "\n\n".join(s.text for s in sections)
            word_count = len(raw_text.split())
            documents.append(
                ParsedDocument(
                    title=book_title,
                    format="txt",
                    pages=0,
                    word_count=word_count,
                    sections=sections,
                    raw_text=raw_text,
                )
            )

        logger.info("parse_kindle_clippings: found %d books", len(documents))
        return documents
