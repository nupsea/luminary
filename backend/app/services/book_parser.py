"""
book_parser.py
==============
Chapter-aware parsing for book-format documents (TXT, PDF, DOCX, MD).

Supports 8 heading pattern families discovered across a corpus of 12 diverse books:
  P1  CHAPTER N. [optional subtitle on next line]     – Alice, Gulliver
  P2  Roman numeral alone  +  subtitle on next line   – Time Machine (Wells)
  P3  I. CAPS TITLE on one line                       – Sherlock Holmes
  P4  Centred roman numeral (leading whitespace ≥4)   – Great Gatsby
  P5  Chapter N: / Chapter N.  arabic/inline          – tech books, Art of Unix
  P6  PART N. / BOOK N. / ordinal BOOK (superstructure) – Gulliver, Marcus Aurelius
  P7  CHAPTER N  (no dot) + confirmed by HERE ENDETH  – Bhagavad Gita
  P8  The [Nth] Book of / The Gospel According        – Bible

Metadata stripping:
  - Gutenberg preamble  (everything before *** START OF …***)
  - Gutenberg footer    (everything after  *** END OF …***)
  - Table-of-Contents blocks  (dense cluster of short heading-like lines near start)
"""

from __future__ import annotations

import logging
import re
from collections import Counter
from pathlib import Path

import chardet
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from markdown_it import MarkdownIt

from app.types import ParsedDocument, Section

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Shared regex patterns
# ---------------------------------------------------------------------------

# Gutenberg delimiters
_RE_PG_START = re.compile(
    r"\*{3}\s*START OF (?:THE |THIS )?PROJECT GUTENBERG EBOOK[^\n]*\*{3}",
    re.IGNORECASE,
)
_RE_PG_END = re.compile(
    r"\*{3}\s*END OF (?:THE |THIS )?PROJECT GUTENBERG EBOOK[^\n]*\*{3}",
    re.IGNORECASE,
)

# Roman numeral helper (I–MMMCMXCIX)
_ROMAN = r"(?:M{0,4}(?:CM|CD|D?C{0,3})(?:XC|XL|L?X{0,3})(?:IX|IV|V?I{1,3}))"

# Ordinal words for "FIRST BOOK", "SECOND BOOK" etc.
_ORDINALS = (
    r"(?:FIRST|SECOND|THIRD|FOURTH|FIFTH|SIXTH|SEVENTH|EIGHTH|NINTH|TENTH|"
    r"ELEVENTH|TWELFTH|THIRTEENTH)"
)

# Pattern families (compiled)
_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    # P1: CHAPTER N. [dot optional] – most explicit
    # Allow optional leading whitespace for indented CHAPTER headings (Gita, etc.)
    (
        "P1",
        re.compile(
            rf"^[ \t]*(?:CHAPTER|CHAP\.?)\s+({_ROMAN}|\d+)[\.:]?\s*$",
            re.MULTILINE | re.IGNORECASE,
        ),
    ),
    # P7: CHAPTER N (no dot, Gita style) – requires HERE ENDETH confirmation
    # Allow leading whitespace
    (
        "P7",
        re.compile(
            rf"^[ \t]*CHAPTER\s+({_ROMAN}|\d+)\s*$",
            re.MULTILINE,
        ),
    ),
    # P3: I. CAPS TITLE on one line (Sherlock Holmes)
    # Require title part to be ALL CAPS, max 80 chars, no lowercase letters
    (
        "P3",
        re.compile(
            rf"^({_ROMAN})\.[ \t]+([A-Z][A-Z ,\'\'\-:/&]{2,79})$",
            re.MULTILINE,
        ),
    ),
    # P6a: PART N. / BOOK N. (superstructure)
    (
        "P6a",
        re.compile(
            rf"^(?:PART|BOOK|SECTION)\s+({_ROMAN}|\d+)[\.]?(?:\s+[^\n]+)?$",
            re.MULTILINE | re.IGNORECASE,
        ),
    ),
    # P6b: ordinal BOOK / ordinal MEDITATION (Marcus Aurelius)
    # Allow leading whitespace
    (
        "P6b",
        re.compile(
            rf"^[ \t]*(?:THE\s+)?{_ORDINALS}\s+BOOK(?:\s+OF\s+[^\n]+)?$",
            re.MULTILINE,
        ),
    ),
    # P8: Bible book headings – match full heading line
    (
        "P8",
        re.compile(
            r"^(?:The\s+(?:First|Second|Third|Fourth|Fifth)\s+Book\s+of\b[^\n]*"
            r"|The\s+Book\s+of\s+\w+[^\n]*"
            r"|The\s+Gospel\s+According\s+to\b[^\n]*"
            r"|The\s+(?:First|Second|Third)\s+Epistle\b[^\n]*"
            r"|The\s+Acts\s+of[^\n]*|The\s+Revelation\b[^\n]*)",
            re.MULTILINE,
        ),
    ),
    # P2: standalone roman numeral line + subtitle on next non-empty line (Wells)
    # Allow leading whitespace for indented chapter numerals
    (
        "P2",
        re.compile(
            rf"^[ \t]*({_ROMAN})\.\s*$",
            re.MULTILINE,
        ),
    ),
    # P4: centred roman numeral (≥4 leading spaces, Gatsby-style)
    (
        "P4",
        re.compile(
            rf"^[ \t]{{4,}}({_ROMAN})[ \t]*$",
            re.MULTILINE | re.IGNORECASE,
        ),
    ),
    # P5: Chapter N: / Chapter N. with inline title (tech books)
    (
        "P5",
        re.compile(
            r"^Chapter\s+(\d+)[\.:\s]\s*(\S.{0,120})$",
            re.MULTILINE,
        ),
    ),
    # Horizontal-rule divider (8+ dashes, equals, or asterisks on own line)
    (
        "HR",
        re.compile(
            r"^[-=*_]{8,}\s*$",
            re.MULTILINE,
        ),
    ),
]

# TOC detection: line that looks like a TOC entry
# Handles: CHAPTER I, I. Title, I Introduction, and space-indented variants
# like ' V In the Golden Age'
_RE_TOC_LINE = re.compile(
    rf"^\s*(?:{_ROMAN}|Chapter\s+\d+|CHAPTER\s+{_ROMAN})"
    r"(?:[.:]\s+\S|\s+[A-Z]\S)",
    re.MULTILINE,
)

# HERE ENDETH confirmation (Gita)
_RE_ENDETH = re.compile(r"HERE ENDETH CHAPTER", re.IGNORECASE)

# Metadata signals that should not become section headings
_METADATA_SIGNALS = frozenset(
    [
        "contents",
        "table of contents",
        "index",
        "preface",
        "foreword",
        "acknowledgements",
        "acknowledgments",
        "copyright",
        "dedication",
        "about the author",
        "bibliography",
        "references",
        "appendix",
        "produced by",
        "project gutenberg",
        "end of the project gutenberg",
    ]
)

# Min chapters required to accept a pattern hit
_MIN_CHAPTERS = 2


# ---------------------------------------------------------------------------
# BookParser
# ---------------------------------------------------------------------------


class BookParser:
    """
    Chapter-aware parser for book-format documents.

    Usage::

        bp = BookParser()
        result = bp.parse(file_path, format="txt")   # returns ParsedDocument or None
        if result is None:
            # document is not book-shaped; use fallback parser
    """

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def parse(self, file_path: Path, fmt: str) -> ParsedDocument | None:
        """
        Attempt to parse *file_path* as a book with chapter structure.

        Returns a :class:`ParsedDocument` when chapter structure is detected,
        or ``None`` when the document doesn't look book-shaped (caller should
        fall back to the generic parser).
        """
        fmt = fmt.lower()
        try:
            if fmt == "txt":
                return self._parse_txt(file_path)
            elif fmt == "pdf":
                return self._parse_pdf(file_path)
            elif fmt in ("docx", "doc"):
                return self._parse_docx(file_path)
            elif fmt in ("md", "markdown"):
                return self._parse_md(file_path)
        except Exception:
            logger.exception("BookParser failed for %s (format=%s)", file_path, fmt)
        return None

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def _parse_txt(self, file_path: Path) -> ParsedDocument | None:
        raw_bytes = file_path.read_bytes()
        detected = chardet.detect(raw_bytes)
        encoding = detected.get("encoding") or "utf-8"
        raw_text = raw_bytes.decode(encoding, errors="replace")

        clean_text, raw_text_no_meta = self._strip_gutenberg(raw_text)
        sections = self._segment_chapters(clean_text)
        if sections is None:
            return None

        return ParsedDocument(
            title=file_path.stem.replace("_", " ").title(),
            format="txt",
            pages=0,
            word_count=len(raw_text_no_meta.split()),
            sections=sections,
            raw_text=raw_text_no_meta,
        )

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _parse_pdf(self, file_path: Path) -> ParsedDocument | None:
        doc = fitz.open(str(file_path))
        if len(doc) == 0:
            return None

        # --- font-size statistics for heading threshold ---
        all_sizes: list[float] = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:  # type: ignore[arg-type]
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            if span["text"].strip():
                                all_sizes.append(span["size"])

        body_avg = sum(all_sizes) / len(all_sizes) if all_sizes else 12.0
        heading_threshold = body_avg * 1.2

        # --- detect running headers/footers by vertical position ---
        # Collect all line y-positions and texts per page
        page_line_texts: list[list[tuple[float, str]]] = []
        for page in doc:
            lines_on_page: list[tuple[float, str]] = []
            for block in page.get_text("dict")["blocks"]:  # type: ignore[arg-type]
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    y = round(line["bbox"][1], -1)  # bucket to nearest 10pt
                    text = " ".join(s["text"] for s in line.get("spans", [])).strip()
                    if text:
                        lines_on_page.append((y, text))
            page_line_texts.append(lines_on_page)

        # Lines appearing verbatim on ≥50% of pages → running header/footer
        text_page_count: Counter[str] = Counter()
        for page_lines in page_line_texts:
            seen = set()
            for _y, t in page_lines:
                if t not in seen:
                    text_page_count[t] += 1
                    seen.add(t)

        n_pages = len(doc)
        running_texts: set[str] = {
            t for t, c in text_page_count.items() if c >= max(2, n_pages * 0.4)
        }

        # --- extract full text, skipping running headers/TOC pages ---
        # Build per-page text to detect TOC pages
        page_texts: list[str] = []
        for _i, (page, page_lines) in enumerate(zip(doc, page_line_texts)):
            lines = [t for _y, t in page_lines if t not in running_texts]
            page_texts.append("\n".join(lines))

        # TOC page: >60% of non-empty lines match TOC pattern
        toc_pages: set[int] = set()
        for pi, ptxt in enumerate(page_texts):
            non_empty = [ln for ln in ptxt.splitlines() if ln.strip()]
            if not non_empty:
                continue
            toc_hits = sum(1 for ln in non_empty if _RE_TOC_LINE.match(ln))
            if toc_hits / len(non_empty) > 0.6:
                toc_pages.add(pi)

        full_text = "\n".join(
            txt for pi, txt in enumerate(page_texts) if pi not in toc_pages
        )

        # --- attempt chapter segmentation on extracted text ---
        sections = self._segment_chapters(full_text)

        if sections is None:
            # Fallback: use font-size headings as sections
            sections = self._pdf_font_sections(doc, running_texts, heading_threshold)

        if not sections:
            return None

        return ParsedDocument(
            title=file_path.stem.replace("_", " ").title(),
            format="pdf",
            pages=n_pages,
            word_count=len(full_text.split()),
            sections=sections,
            raw_text=full_text,
        )

    def _pdf_font_sections(
        self,
        doc: fitz.Document,
        running_texts: set[str],
        heading_threshold: float,
    ) -> list[Section]:
        """Fallback: segment PDF by font-size heading heuristic."""
        sections: list[Section] = []
        current_heading = "Introduction"
        current_level = 1
        current_page_start = 0
        current_texts: list[str] = []

        def flush(next_heading: str, next_level: int, next_page: int) -> None:
            nonlocal current_heading, current_level, current_page_start, current_texts
            text = "\n".join(current_texts).strip()
            if text and not _is_metadata(current_heading):
                sections.append(
                    Section(
                        heading=current_heading,
                        level=current_level,
                        text=text,
                        page_start=current_page_start,
                        page_end=next_page - 1,
                    )
                )
            current_heading, current_level, current_page_start, current_texts = (
                next_heading,
                next_level,
                next_page,
                [],
            )

        for page_num, page in enumerate(doc):
            for block in page.get_text("dict")["blocks"]:  # type: ignore[arg-type]
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    max_size = max(s["size"] for s in spans)
                    line_text = " ".join(s["text"] for s in spans).strip()
                    if not line_text or line_text in running_texts:
                        continue
                    if max_size >= heading_threshold and len(line_text) < 120:
                        flush(line_text, 1, page_num + 1)
                    else:
                        current_texts.append(line_text)

        flush("_end", 0, len(doc))
        return sections

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _parse_docx(self, file_path: Path) -> ParsedDocument | None:
        doc = DocxDocument(str(file_path))
        sections: list[Section] = []
        raw_parts: list[str] = []
        current_heading = "Introduction"
        current_level = 1
        current_texts: list[str] = []

        def flush(next_heading: str, next_level: int) -> None:
            nonlocal current_heading, current_level, current_texts
            text = "\n".join(current_texts).strip()
            if text and not _is_metadata(current_heading):
                sections.append(
                    Section(
                        heading=current_heading,
                        level=current_level,
                        text=text,
                        page_start=0,
                        page_end=0,
                    )
                )
            current_heading, current_level, current_texts = next_heading, next_level, []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            # Word built-in heading styles
            m_style = re.match(r"Heading\s+(\d+)", style_name)
            if m_style:
                flush(text, int(m_style.group(1)))
                continue

            # Regex fallback on paragraph text (catches books saved without styles)
            m_chapter = _match_any_pattern(text)
            if m_chapter:
                flush(text, m_chapter)
                continue

            current_texts.append(text)
            raw_parts.append(text)

        flush("_end", 0)

        if len(sections) < _MIN_CHAPTERS:
            return None

        raw_text = "\n".join(raw_parts)
        return ParsedDocument(
            title=file_path.stem.replace("_", " ").title(),
            format="docx",
            pages=0,
            word_count=len(raw_text.split()),
            sections=sections,
            raw_text=raw_text,
        )

    # ------------------------------------------------------------------
    # Markdown
    # ------------------------------------------------------------------

    def _parse_md(self, file_path: Path) -> ParsedDocument | None:
        raw_text = file_path.read_text(encoding="utf-8", errors="replace")
        md = MarkdownIt()
        tokens = md.parse(raw_text)

        sections: list[Section] = []
        current_heading = "Introduction"
        current_level = 1
        current_texts: list[str] = []
        in_heading = False
        pending_heading = ""
        pending_level = 1

        def flush(next_heading: str, next_level: int) -> None:
            nonlocal current_heading, current_level, current_texts
            body = "\n".join(current_texts).strip()
            if body and not _is_metadata(current_heading):
                sections.append(
                    Section(
                        heading=current_heading,
                        level=current_level,
                        text=body,
                        page_start=0,
                        page_end=0,
                    )
                )
            current_heading, current_level, current_texts = next_heading, next_level, []

        for token in tokens:
            if token.type == "heading_open":
                in_heading = True
                pending_level = int(token.tag[1]) if token.tag else 1
            elif token.type == "inline" and in_heading:
                pending_heading = token.content
            elif token.type == "heading_close":
                in_heading = False
                flush(pending_heading, pending_level)
            elif token.type == "inline":
                current_texts.append(token.content)

        flush("_end", 0)

        if len(sections) < _MIN_CHAPTERS:
            return None

        return ParsedDocument(
            title=file_path.stem.replace("_", " ").title(),
            format="md",
            pages=0,
            word_count=len(raw_text.split()),
            sections=sections,
            raw_text=raw_text,
        )

    # ------------------------------------------------------------------
    # Core: Gutenberg stripping + chapter segmentation (used by TXT + PDF)
    # ------------------------------------------------------------------

    def _strip_gutenberg(self, text: str) -> tuple[str, str]:
        """
        Returns (clean_content_for_segmentation, raw_text_without_legal_boilerplate).

        Strips Gutenberg preamble, footer, and TOC block.
        Also normalises CRLF → LF so all multiline patterns work uniformly
        regardless of the source file's line endings.
        """
        # Normalise Windows CRLF to Unix LF – must happen before any regex matching
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Strip preamble
        m_start = _RE_PG_START.search(text)
        if m_start:
            text = text[m_start.end():]

        # Strip footer
        m_end = _RE_PG_END.search(text)
        if m_end:
            text = text[: m_end.start()]

        raw_clean = text  # keep for word_count / raw_text

        # Strip TOC block: dense cluster of short TOC-like lines in first 20%
        text = self._strip_toc_block(text)

        return text, raw_clean

    def _strip_toc_block(self, text: str) -> str:
        """Remove TOC block if it exists in the first 20% of content."""
        cutoff = max(len(text) // 5, 5000)
        head = text[:cutoff]
        lines = head.splitlines()

        toc_start = toc_end = -1
        run = 0
        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                if run >= 3:
                    toc_end = i
                run = 0
                continue
            if _RE_TOC_LINE.match(line) or re.match(
                r"^\s*(?:Contents|TABLE OF CONTENTS)\s*$", line, re.IGNORECASE
            ):
                if toc_start == -1:
                    toc_start = i
                run += 1
                toc_end = i
            else:
                if run >= 3:
                    break  # end of TOC
                toc_start = toc_end = -1
                run = 0

        if toc_start != -1 and toc_end != -1 and toc_end - toc_start >= 3:
            kept_before = "\n".join(lines[:toc_start])
            kept_after_start = toc_end + 1
            # skip blank lines after TOC
            while kept_after_start < len(lines) and not lines[kept_after_start].strip():
                kept_after_start += 1
            rest_head = "\n".join(lines[kept_after_start:])
            text = kept_before + "\n" + rest_head + text[cutoff:]

        return text

    def _segment_chapters(self, text: str) -> list[Section] | None:
        """
        Detect which heading pattern dominates and split text into Sections.

        Returns None when no pattern yields ≥ _MIN_CHAPTERS matches.
        """
        best_pattern_id, best_pattern, match_count = self._pick_best_pattern(text)
        if best_pattern_id is None or match_count < _MIN_CHAPTERS:
            return None

        logger.debug(
            "BookParser: using pattern %s (%d matches)", best_pattern_id, match_count
        )

        if best_pattern_id in ("P2", "P4"):
            return self._split_roman_with_subtitle(text, best_pattern, best_pattern_id)
        elif best_pattern_id == "P1":
            return self._split_with_optional_subtitle(text, best_pattern, "CHAPTER")
        elif best_pattern_id == "HR":
            return self._split_by_hr(text)
        else:
            return self._split_generic(text, best_pattern)

    def _pick_best_pattern(
        self, text: str
    ) -> tuple[str | None, re.Pattern[str] | None, int]:
        """Return (pattern_id, pattern, count) for the best-matching pattern."""
        # P7 requires HERE ENDETH confirmation
        has_endeth = bool(_RE_ENDETH.search(text))

        best_id: str | None = None
        best_pat: re.Pattern[str] | None = None
        best_count = 0

        for pid, pat in _PATTERNS:
            if pid == "P7" and not has_endeth:
                continue
            matches = pat.findall(text)
            count = len(matches)
            if count > best_count:
                best_count = count
                best_id = pid
                best_pat = pat

        return best_id, best_pat, best_count

    # ------------------------------------------------------------------
    # Splitting helpers
    # ------------------------------------------------------------------

    def _split_generic(
        self, text: str, pattern: re.Pattern[str]
    ) -> list[Section] | None:
        """Generic split: each match becomes a heading; text between = body."""
        sections: list[Section] = []
        positions = [(m.start(), m.end(), m.group(0).strip()) for m in pattern.finditer(text)]

        if len(positions) < _MIN_CHAPTERS:
            return None

        for i, (start, end, heading) in enumerate(positions):
            if _is_metadata(heading):
                continue
            body_start = end
            body_end = positions[i + 1][0] if i + 1 < len(positions) else len(text)
            body = text[body_start:body_end].strip()
            # For P1/P7 the actual subtitle may appear on the next non-empty line
            sections.append(
                Section(
                    heading=_clean_heading(heading),
                    level=1,
                    text=body,
                    page_start=0,
                    page_end=0,
                )
            )

        return sections if len(sections) >= _MIN_CHAPTERS else None

    def _split_with_optional_subtitle(
        self, text: str, pattern: re.Pattern[str], prefix: str
    ) -> list[Section] | None:
        """
        For P1-style patterns: the heading line is e.g. 'CHAPTER I.'
        The next non-empty line may be a short subtitle (< 80 chars, not a sentence).
        """
        sections: list[Section] = []
        matches = list(pattern.finditer(text))
        if len(matches) < _MIN_CHAPTERS:
            return None

        for i, m in enumerate(matches):
            heading_line = m.group(0).strip()
            body_start = m.end()
            body_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body_chunk = text[body_start:body_end]

            # Check for subtitle on next non-empty line
            remaining_lines = body_chunk.lstrip("\r\n").splitlines()
            subtitle = ""
            body_line_start = 0
            for li, line in enumerate(remaining_lines):
                stripped = line.strip()
                if stripped:
                    # subtitle: short ≤ 80 chars, not ending with sentence punctuation
                    if len(stripped) <= 80 and stripped[-1] not in ".?!":
                        subtitle = stripped
                        body_line_start = li + 1
                    break

            full_heading = (
                f"{heading_line} — {subtitle}" if subtitle else heading_line
            )
            body = "\n".join(remaining_lines[body_line_start:]).strip()

            if _is_metadata(full_heading):
                continue

            sections.append(
                Section(
                    heading=_clean_heading(full_heading),
                    level=1,
                    text=body,
                    page_start=0,
                    page_end=0,
                )
            )

        return sections if len(sections) >= _MIN_CHAPTERS else None

    def _split_roman_with_subtitle(
        self, text: str, pattern: re.Pattern[str], pid: str
    ) -> list[Section] | None:
        """
        For P2 (roman. on own line) and P4 (centred roman):
        the next non-empty line is always the section title.
        """
        sections: list[Section] = []
        matches = list(pattern.finditer(text))
        if len(matches) < _MIN_CHAPTERS:
            return None

        for i, m in enumerate(matches):
            body_start = m.end()
            body_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body_chunk = text[body_start:body_end].lstrip("\r\n")

            lines = body_chunk.splitlines()
            subtitle = ""
            body_line_start = 0
            for li, line in enumerate(lines):
                stripped = line.strip()
                if stripped:
                    subtitle = stripped
                    body_line_start = li + 1
                    break

            numeral = m.group(1) if m.lastindex else m.group(0).strip()
            heading = f"{numeral}. {subtitle}" if subtitle else numeral
            body = "\n".join(lines[body_line_start:]).strip()

            if _is_metadata(heading):
                continue

            sections.append(
                Section(
                    heading=_clean_heading(heading),
                    level=1,
                    text=body,
                    page_start=0,
                    page_end=0,
                )
            )

        return sections if len(sections) >= _MIN_CHAPTERS else None

    def _split_by_hr(self, text: str) -> list[Section] | None:
        """Split on horizontal-rule dividers (-----)."""
        hr_pat = _PATTERNS[-1][1]  # HR is last
        parts = hr_pat.split(text)
        sections: list[Section] = []
        for i, part in enumerate(parts):
            body = part.strip()
            if not body:
                continue
            # First non-empty line of each part becomes the heading
            lines = body.splitlines()
            first_line = next((ln.strip() for ln in lines if ln.strip()), f"Section {i+1}")
            rest = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
            if not rest:
                rest = body
            heading = first_line if len(first_line) <= 120 else f"Section {i+1}"
            sections.append(
                Section(
                    heading=_clean_heading(heading),
                    level=1,
                    text=rest or body,
                    page_start=0,
                    page_end=0,
                )
            )
        return sections if len(sections) >= _MIN_CHAPTERS else None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _is_metadata(heading: str) -> bool:
    """Return True if heading looks like book metadata (not a chapter)."""
    return heading.lower().strip() in _METADATA_SIGNALS


def _clean_heading(heading: str) -> str:
    """Normalise whitespace and trim trailing punctuation artefacts."""
    h = re.sub(r"\s+", " ", heading).strip()
    # Remove trailing lone dots (e.g. "CHAPTER I.")
    h = re.sub(r"\.\s*$", "", h).strip()
    return h


def _match_any_pattern(text: str) -> int | None:
    """
    Try matching a single line of text against known patterns.
    Returns heading level (1 or 2) if matched, else None.
    Used by the DOCX paragraph fallback.
    """
    for pid, pat in _PATTERNS:
        if pid == "HR":
            continue
        if pat.match(text.strip()):
            return 2 if pid in ("P6a", "P6b", "P8") else 1
    return None
