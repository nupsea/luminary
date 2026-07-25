import pytest

from app.services.parser import (
    DocumentParser,
    _heading_is_prose,
    _join_spans,
    _norm_ws,
    _sections_plausible,
)
from app.types import ParsedDocument, Section


@pytest.fixture(scope="session")
def parser():
    return DocumentParser()


@pytest.fixture(scope="session")
def tmp_fixtures(tmp_path_factory):
    return tmp_path_factory.mktemp("fixtures")


@pytest.fixture(scope="session")
def pdf_file(tmp_fixtures):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    path = tmp_fixtures / "sample.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    elements = [
        Paragraph("Introduction", styles["Heading1"]),
        Paragraph("This is the introduction text with some content.", styles["Normal"]),
        Paragraph("Methods", styles["Heading1"]),
        Paragraph("This describes the methods used in this study.", styles["Normal"]),
    ]
    doc.build(elements)
    return path


@pytest.fixture(scope="session")
def docx_file(tmp_fixtures):
    from docx import Document

    path = tmp_fixtures / "sample.docx"
    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("This is the content of chapter one.")
    doc.add_heading("Section 1.1", level=2)
    doc.add_paragraph("A subsection with more detail.")
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("Content for chapter two.")
    doc.save(str(path))
    return path


@pytest.fixture(scope="session")
def txt_file(tmp_fixtures):
    path = tmp_fixtures / "sample.txt"
    content = (
        "First paragraph with some text.\n\nSecond paragraph that follows.\n\nThird paragraph here."
    )
    path.write_text(content, encoding="utf-8")
    return path


@pytest.fixture(scope="session")
def md_file(tmp_fixtures):
    path = tmp_fixtures / "sample.md"
    content = (
        "# Title\n\nIntroduction text.\n\n"
        "## Background\n\nBackground content.\n\n"
        "## Methods\n\nMethods content."
    )
    path.write_text(content, encoding="utf-8")
    return path


def test_pdf_parse_returns_sections(parser, pdf_file):
    result = parser.parse(pdf_file, "pdf")
    assert isinstance(result, ParsedDocument)
    assert result.format == "pdf"
    assert result.pages > 0
    assert result.word_count > 0
    assert len(result.sections) >= 1


def test_docx_parse_returns_heading_hierarchy(parser, docx_file):
    result = parser.parse(docx_file, "docx")
    assert isinstance(result, ParsedDocument)
    assert result.format == "docx"
    headings = [s.heading for s in result.sections]
    assert "Chapter One" in headings or len(headings) >= 1
    levels = [s.level for s in result.sections]
    assert any(lvl == 1 for lvl in levels)


def test_txt_parse_detects_encoding_and_sections(parser, txt_file):
    result = parser.parse(txt_file, "txt")
    assert isinstance(result, ParsedDocument)
    assert result.format == "txt"
    assert len(result.sections) >= 2
    assert result.word_count > 0


def test_md_parse_detects_headings(parser, md_file):
    result = parser.parse(md_file, "md")
    assert isinstance(result, ParsedDocument)
    assert result.format == "md"
    headings = [s.heading for s in result.sections]
    assert "Title" in headings or "Background" in headings or "Methods" in headings


@pytest.mark.parametrize(
    "heading,expected",
    [
        ("Chapter 5 describes the foundations of machine learning", True),
        ("Chapter 7 focuses on modern developments", True),
        ("Chapter 11 along with some of their fundamental properties", True),
        ("Part 3 are the same triangle", True),
        ("volume 22 of Carus Mathematical Monographs", True),
        ("Chapter 5. Machine Learning", False),
        ("Chapter 5: Optimization", False),
        ("Chapter 5 Machine Learning", False),
        ("Chapter 1 — The Beginning", False),
        ("CHAPTER I", False),
        ("Introduction", False),
        ("Methods", False),
    ],
)
def test_heading_is_prose(heading, expected):
    assert _heading_is_prose(heading) is expected


def test_sections_plausible_rejects_all_prose():
    secs = [
        Section(heading=h, level=1, text="x", page_start=1, page_end=1)
        for h in (
            "Chapter 5 describes the foundations",
            "Chapter 7 focuses on developments",
            "Part 3 are the same triangle",
        )
    ]
    assert _sections_plausible(secs) is False


def test_sections_plausible_accepts_real_headings():
    secs = [
        Section(heading=h, level=1, text="x", page_start=1, page_end=1)
        for h in ("Chapter 1 — The Beginning", "Chapter 2 — The Middle", "Chapter 3 — The End")
    ]
    assert _sections_plausible(secs) is True


def test_sections_plausible_empty():
    assert _sections_plausible([]) is False


@pytest.fixture(scope="session")
def prose_trap_pdf(tmp_fixtures):
    """A PDF whose preface enumerates its own chapters in body font — the exact
    trap that made chapter-regex matching produce prose-fragment headings."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    path = tmp_fixtures / "prose_trap.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    elements = [
        Paragraph("Preface", styles["Heading1"]),
        Paragraph(
            "Chapter 5 describes the foundations of machine learning, both algorithms "
            "for optimizing over given training examples and their analysis.",
            styles["Normal"],
        ),
        Paragraph(
            "Chapter 7 focuses on modern developments in understanding these methods.",
            styles["Normal"],
        ),
        Paragraph(
            "Part 3 discusses ranking and social choice as well as related problems.",
            styles["Normal"],
        ),
        Paragraph("Introduction", styles["Heading1"]),
        Paragraph(
            "The real introduction body with enough content to fill a section.",
            styles["Normal"],
        ),
        Paragraph("Methods", styles["Heading1"]),
        Paragraph(
            "This describes the actual methods used in this study in detail.",
            styles["Normal"],
        ),
    ]
    doc.build(elements)
    return path


def test_pdf_preface_prose_not_mistaken_for_headings(parser, prose_trap_pdf):
    result = parser.parse(prose_trap_pdf, "pdf")
    headings = [s.heading for s in result.sections]
    assert not any(_heading_is_prose(h) for h in headings), headings
    assert not any(h.lower().startswith("chapter 5 describes") for h in headings), headings
    # Headings carry no multi-space runs (PyMuPDF span-join artifact).
    assert not any("  " in h for h in headings), headings
    # The real font-based headings survive.
    joined = " ".join(headings).lower()
    assert any(k in joined for k in ("introduction", "methods", "preface")), headings


@pytest.mark.parametrize(
    "raw,expected",
    [
        ("The   Law   of   Large   Numbers", "The Law of Large Numbers"),
        ("  Best Rank- k   Approximations ", "Best Rank- k Approximations"),
        ("Singular\tValue\nDecomposition", "Singular Value Decomposition"),
        ("Introduction", "Introduction"),
    ],
)
def test_norm_ws(raw, expected):
    assert _norm_ws(raw) == expected


def _span(text: str, x0: float, x1: float, size: float = 12.0) -> dict:
    return {"text": text, "bbox": (x0, 0.0, x1, size), "size": size}


class TestJoinSpans:
    def test_glyph_split_word_is_not_broken_by_spaces(self):
        """A PDF alternating Type3 font subsets splits one word across many spans
        with no horizontal gap; joining on a space corrupted every word."""
        spans = [
            _span("The C", 27.750, 74.964, 16.5),
            _span("on", 74.980, 95.581, 16.5),
            _span("ce", 95.424, 113.169, 16.5),
            _span("ptu", 113.177, 141.037, 16.5),
            _span("al F", 141.045, 169.697, 16.5),
            _span("oun", 169.028, 200.460, 16.5),
            _span("da", 200.303, 220.153, 16.5),
            _span("t", 219.748, 226.282, 16.5),
            _span("i", 226.299, 231.562, 16.5),
            _span("ons", 231.570, 260.296, 16.5),
        ]
        assert _join_spans(spans) == "The Conceptual Foundations"

    def test_negative_kerning_does_not_insert_a_space(self):
        assert _join_spans([_span("Ke", 10.0, 20.0), _span("rn", 19.4, 30.0)]) == "Kern"

    def test_positional_word_gap_becomes_a_space(self):
        """Words separated by position rather than a space glyph must not weld."""
        spans = [_span("Hello", 10.0, 40.0), _span("world", 44.0, 74.0)]
        assert _join_spans(spans) == "Hello world"

    def test_existing_space_is_not_doubled(self):
        spans = [_span("Hello ", 10.0, 43.0), _span("world", 44.0, 74.0)]
        assert _join_spans(spans) == "Hello world"

    def test_leading_space_on_next_span_is_not_doubled(self):
        spans = [_span("Hello", 10.0, 40.0), _span(" world", 44.0, 78.0)]
        assert _join_spans(spans) == "Hello world"

    def test_empty_spans_are_skipped(self):
        assert _join_spans([_span("a", 10.0, 15.0), _span("", 15.0, 15.0)]) == "a"

    def test_no_spans(self):
        assert _join_spans([]) == ""

    def test_gap_threshold_scales_with_font_size(self):
        """A 3pt gap is a word break at 12pt but mere kerning at 24pt."""
        assert _join_spans([_span("a", 0.0, 10.0, 12.0), _span("b", 13.0, 20.0, 12.0)]) == "a b"
        assert _join_spans([_span("a", 0.0, 10.0, 24.0), _span("b", 13.0, 20.0, 24.0)]) == "ab"
